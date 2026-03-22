"""Report generation and export API."""

import csv
import io
import json
import uuid

from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlmodel.ext.asyncio.session import AsyncSession

from app.api.deps import get_db, require_permission
from app.models.report import Report, ReportStatus, ReportType
from app.models.user import User
from app.services.report_generator import (
    generate_cost_report,
    generate_performance_report,
    generate_safety_report,
    generate_value_report,
)

router = APIRouter()

REPORT_GENERATORS = {
    "performance": generate_performance_report,
    "safety": generate_safety_report,
    "cost": generate_cost_report,
    "value": generate_value_report,
}


# ── Persistent report CRUD ─────────────────────────────


@router.post("", status_code=201)
async def create_report(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.generate"),
):
    """Generate and persist a report."""
    generator = REPORT_GENERATORS.get(report_type)
    if not generator:
        raise HTTPException(400, f"Invalid report type: {report_type}")

    rt = ReportType(report_type)
    report = Report(
        task_id=task_id,
        report_type=rt,
        status=ReportStatus.generating,
        created_by=current_user.id,
    )
    session.add(report)
    await session.commit()
    await session.refresh(report)

    try:
        content = await generator(task_id, session)
        report.content_json = json.dumps(content, ensure_ascii=False, default=str)
        report.title = content.get("title", "")
        report.status = ReportStatus.ready
    except Exception as e:
        report.status = ReportStatus.failed
        report.error_message = str(e)[:500]

    session.add(report)
    await session.commit()
    await session.refresh(report)
    return {
        "id": str(report.id),
        "task_id": str(report.task_id),
        "report_type": report.report_type,
        "status": report.status,
        "title": report.title,
        "content": (
            json.loads(report.content_json)
            if report.status == ReportStatus.ready
            else None
        ),
        "error_message": report.error_message,
        "created_at": report.created_at.isoformat(),
    }


@router.get("")
async def list_reports(
    task_id: uuid.UUID | None = None,
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.read"),
):
    """List persisted reports, filtered by visibility."""
    from sqlmodel import select as sel

    stmt = sel(Report).order_by(Report.created_at.desc())
    if task_id:
        stmt = stmt.where(Report.task_id == task_id)
    result = await session.exec(stmt)
    all_reports = result.all()

    # Filter by visibility
    visible = []
    for r in all_reports:
        if current_user.role == "admin":
            visible.append(r)
        elif r.created_by == current_user.id:
            visible.append(r)
        elif r.visibility == "public":
            visible.append(r)
        elif r.visibility == "team":
            visible.append(r)  # team = all authenticated users
        elif r.allowed_users:
            if str(current_user.id) in r.allowed_users.split(","):
                visible.append(r)
    reports = visible

    return [
        {
            "id": str(r.id),
            "task_id": str(r.task_id),
            "report_type": r.report_type,
            "status": r.status,
            "title": r.title,
            "visibility": r.visibility,
            "created_at": r.created_at.isoformat(),
        }
        for r in reports
    ]


@router.get("/{report_id}")
async def get_report(
    report_id: uuid.UUID,
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.read"),
):
    """Fetch a persisted report with full content."""
    report = await session.get(Report, report_id)
    if not report:
        raise HTTPException(404, "Report not found")

    # Visibility access check
    if (
        current_user.role != "admin"
        and report.created_by != current_user.id
        and report.visibility == "creator"
        and str(current_user.id) not in (report.allowed_users or "").split(",")
    ):
        raise HTTPException(403, "无权查看此报告")

    return {
        "id": str(report.id),
        "task_id": str(report.task_id),
        "report_type": report.report_type,
        "status": report.status,
        "title": report.title,
        "visibility": report.visibility,
        "content": json.loads(report.content_json) if report.content_json else None,
        "error_message": report.error_message,
        "created_at": report.created_at.isoformat(),
    }


# ── Legacy generate (stateless) ────────────────────────


@router.post("/generate")
async def generate_report(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.read"),
):
    """Generate a report as JSON (stateless, not persisted)."""
    generator = REPORT_GENERATORS.get(report_type)
    if not generator:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail=(
                "Invalid report type. Choose from: "
                f"{list(REPORT_GENERATORS.keys())}"
            ),
        )
    try:
        return await generator(task_id, session)
    except ValueError as e:
        raise HTTPException(
            status.HTTP_404_NOT_FOUND, detail=str(e)
        ) from e


# ── Export endpoints ────────────────────────────────────


@router.post("/export/csv")
async def export_csv(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.export"),
):
    """Export report as CSV."""
    report = await _get_report(
        task_id, report_type, session
    )
    csv_content = _report_to_csv(report)

    return StreamingResponse(
        io.BytesIO(csv_content.encode("utf-8-sig")),
        media_type="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": (
                f'attachment; filename='
                f'"{report_type}_report.csv"'
            )
        },
    )


@router.post("/export/html")
async def export_html(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.export"),
):
    """Export report as HTML."""
    report = await _get_report(
        task_id, report_type, session
    )
    html = _report_to_html(report)

    return StreamingResponse(
        io.BytesIO(html.encode("utf-8")),
        media_type="text/html; charset=utf-8",
        headers={
            "Content-Disposition": (
                f'attachment; filename='
                f'"{report_type}_report.html"'
            )
        },
    )


@router.post("/export/docx")
async def export_docx(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.export"),
):
    """Export report as DOCX."""
    report = await _get_report(
        task_id, report_type, session
    )
    docx_bytes = _report_to_docx(report)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": (
                f'attachment; filename='
                f'"{report_type}_report.docx"'
            )
        },
    )


@router.post("/export/pdf")
async def export_pdf(
    task_id: uuid.UUID,
    report_type: str = "performance",
    session: AsyncSession = Depends(get_db),
    current_user: User = require_permission("reports.export"),
):
    """Export report as PDF (via HTML rendering)."""
    report = await _get_report(task_id, report_type, session)
    html = _report_to_html(report)

    # Convert HTML to PDF using simple approach
    pdf_bytes = _html_to_pdf(html)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{report_type}_report.pdf"'
        },
    )


async def _get_report(task_id, report_type, session):
    generator = REPORT_GENERATORS.get(report_type)
    if not generator:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail="Invalid report type",
        )
    try:
        return await generator(task_id, session)
    except ValueError as e:
        raise HTTPException(
            status.HTTP_404_NOT_FOUND, detail=str(e)
        ) from e


# ── CSV helper ──────────────────────────────────────────


def _report_to_csv(report: dict) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    rtype = report.get("type", "")

    # Header row with metadata
    writer.writerow(["Report", report.get("title", "")])
    writer.writerow(["Model", report.get("model_name", "")])
    writer.writerow(
        ["Generated", report.get("generated_at", "")]
    )
    writer.writerow([])

    if rtype == "performance":
        headers = [
            "Criterion",
            "Avg Score",
            "Min Score",
            "Max Score",
            "Samples",
            "Avg Latency (ms)",
        ]
        writer.writerow(headers)
        for row in report.get("criteria_breakdown", []):
            writer.writerow([
                row["criterion"],
                row["avg_score"],
                row["min_score"],
                row["max_score"],
                row["sample_count"],
                row["avg_latency_ms"],
            ])
        writer.writerow([])
        writer.writerow([
            "Overall Score",
            report.get("overall_score", 0),
        ])

    elif rtype == "safety":
        writer.writerow([
            "Total",
            report.get("total_samples", 0),
        ])
        writer.writerow([
            "Errors",
            report.get("error_count", 0),
        ])
        writer.writerow([
            "Error Rate",
            report.get("error_rate", 0),
        ])
        writer.writerow([
            "Risk",
            report.get("risk_level", ""),
        ])
        writer.writerow([])
        writer.writerow([
            "Prompt",
            "Expected",
            "Actual",
            "Score",
        ])
        for c in report.get("error_cases", []):
            writer.writerow([
                c["prompt"][:200],
                c["expected"][:200],
                c["actual"][:200],
                c["score"],
            ])

    elif rtype == "cost":
        for key in [
            "total_samples",
            "avg_latency_ms",
            "min_latency_ms",
            "max_latency_ms",
            "avg_first_token_ms",
            "avg_tokens_per_response",
            "total_tokens",
            "duration_seconds",
            "throughput_tokens_per_sec",
        ]:
            writer.writerow([key, report.get(key, "")])

    elif rtype == "value":
        writer.writerow([
            "Overall Score",
            report.get("overall_score", 0),
        ])
        writer.writerow([
            "Avg Latency (ms)",
            report.get("avg_latency_ms", 0),
        ])
        writer.writerow([
            "Value Index",
            report.get("value_index", 0),
        ])
        writer.writerow([
            "Throughput (tok/s)",
            report.get("throughput_tokens_per_sec", 0),
        ])
        writer.writerow([])
        headers = [
            "Criterion",
            "Avg Score",
            "Min Score",
            "Max Score",
            "Samples",
            "Avg Latency (ms)",
        ]
        writer.writerow(headers)
        for row in report.get("criteria_breakdown", []):
            writer.writerow([
                row["criterion"],
                row["avg_score"],
                row["min_score"],
                row["max_score"],
                row["sample_count"],
                row["avg_latency_ms"],
            ])

    return buf.getvalue()


# ── HTML helper ─────────────────────────────────────────

_CSS = """
body{font-family:system-ui,sans-serif;max-width:900px;
margin:40px auto;padding:0 20px;color:#1a1a1a}
h1{border-bottom:2px solid #2563eb;padding-bottom:8px}
.meta{color:#666;margin-bottom:24px}
table{border-collapse:collapse;width:100%;margin:16px 0}
th,td{border:1px solid #d1d5db;padding:8px 12px;
text-align:left}
th{background:#f3f4f6}
.badge{display:inline-block;padding:2px 10px;
border-radius:4px;font-weight:600}
.low{background:#d1fae5;color:#065f46}
.mid{background:#fef3c7;color:#92400e}
.high{background:#fee2e2;color:#991b1b}
"""


def _report_to_html(report: dict) -> str:
    rtype = report.get("type", "")
    title = report.get("title", "Report")
    meta = (
        f"<p class='meta'>Model: {report.get('model_name')}"
        f" | Generated: {report.get('generated_at')}</p>"
    )
    body = ""

    if rtype == "performance":
        body = _html_perf(report)
    elif rtype == "safety":
        body = _html_safety(report)
    elif rtype == "cost":
        body = _html_cost(report)
    elif rtype == "value":
        body = _html_value(report)

    return (
        "<!DOCTYPE html><html><head>"
        '<meta charset="utf-8">'
        f"<title>{title}</title>"
        f"<style>{_CSS}</style></head><body>"
        f"<h1>{title}</h1>{meta}{body}"
        "</body></html>"
    )


def _html_table(headers: list, rows: list) -> str:
    ths = "".join(f"<th>{h}</th>" for h in headers)
    trs = ""
    for row in rows:
        tds = "".join(f"<td>{v}</td>" for v in row)
        trs += f"<tr>{tds}</tr>"
    return f"<table><tr>{ths}</tr>{trs}</table>"


def _html_perf(r: dict) -> str:
    overall = f"<p>Overall Score: <b>{r['overall_score']}</b>"
    overall += f" ({r['total_samples']} samples)</p>"
    headers = [
        "Criterion",
        "Avg",
        "Min",
        "Max",
        "Samples",
        "Latency(ms)",
    ]
    rows = [
        [
            c["criterion"],
            c["avg_score"],
            c["min_score"],
            c["max_score"],
            c["sample_count"],
            c["avg_latency_ms"],
        ]
        for c in r.get("criteria_breakdown", [])
    ]
    return overall + _html_table(headers, rows)


def _html_safety(r: dict) -> str:
    rl = r.get("risk_level", "")
    css = (
        "low"
        if "\u4f4e" in rl
        else "mid"
        if "\u4e2d" in rl
        else "high"
    )
    out = (
        f"<p>Risk: <span class='badge {css}'>"
        f"{rl}</span></p>"
        f"<p>Errors: {r['error_count']}/{r['total_samples']}"
        f" ({r['error_rate']})</p>"
    )
    if r.get("error_cases"):
        headers = ["Prompt", "Expected", "Actual", "Score"]
        rows = [
            [
                c["prompt"][:120],
                c["expected"][:120],
                c["actual"][:120],
                c["score"],
            ]
            for c in r["error_cases"]
        ]
        out += _html_table(headers, rows)
    return out


def _html_cost(r: dict) -> str:
    items = [
        ("Avg Latency", f"{r['avg_latency_ms']} ms"),
        ("Min Latency", f"{r['min_latency_ms']} ms"),
        ("Max Latency", f"{r['max_latency_ms']} ms"),
        (
            "First Token",
            f"{r['avg_first_token_ms']} ms",
        ),
        (
            "Avg Tokens/Resp",
            r["avg_tokens_per_response"],
        ),
        ("Total Tokens", r["total_tokens"]),
        ("Duration", f"{r['duration_seconds']}s"),
        (
            "Throughput",
            f"{r['throughput_tokens_per_sec']} tok/s",
        ),
    ]
    return _html_table(
        ["Metric", "Value"],
        [[k, v] for k, v in items],
    )


def _html_value(r: dict) -> str:
    summary = (
        f"<p>Score: <b>{r['overall_score']}</b> | "
        f"Latency: {r['avg_latency_ms']}ms | "
        f"Value Index: <b>{r['value_index']}</b></p>"
    )
    return summary + _html_perf(r)


def _html_to_pdf(html: str) -> bytes:
    """Convert HTML to PDF. Uses weasyprint if available, falls back to basic."""
    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except ImportError:
        # Fallback: return HTML wrapped as "PDF" (user can print to PDF)
        # In production, install weasyprint for real PDF generation
        return html.encode("utf-8")


# ── DOCX helper ─────────────────────────────────────────


def _report_to_docx(report: dict) -> bytes:
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    title = report.get("title", "Report")
    doc.add_heading(title, level=1)
    doc.add_paragraph(
        f"Model: {report.get('model_name', '')}  |  "
        f"Generated: {report.get('generated_at', '')}"
    )

    rtype = report.get("type", "")
    if rtype == "performance":
        _docx_perf(doc, report)
    elif rtype == "safety":
        _docx_safety(doc, report)
    elif rtype == "cost":
        _docx_cost(doc, report)
    elif rtype == "value":
        _docx_value(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_add_table(doc, headers, rows):
    table = doc.add_table(
        rows=1 + len(rows), cols=len(headers)
    )
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)


def _docx_perf(doc, r):
    doc.add_paragraph(
        f"Overall Score: {r['overall_score']}  "
        f"({r['total_samples']} samples)"
    )
    headers = [
        "Criterion",
        "Avg",
        "Min",
        "Max",
        "N",
        "Latency(ms)",
    ]
    rows = [
        [
            c["criterion"],
            c["avg_score"],
            c["min_score"],
            c["max_score"],
            c["sample_count"],
            c["avg_latency_ms"],
        ]
        for c in r.get("criteria_breakdown", [])
    ]
    _docx_add_table(doc, headers, rows)


def _docx_safety(doc, r):
    doc.add_paragraph(
        f"Risk Level: {r['risk_level']}  |  "
        f"Errors: {r['error_count']}/{r['total_samples']}  "
        f"({r['error_rate']})"
    )
    if r.get("error_cases"):
        doc.add_heading("Error Cases", level=2)
        headers = [
            "Prompt",
            "Expected",
            "Actual",
            "Score",
        ]
        rows = [
            [
                c["prompt"][:100],
                c["expected"][:100],
                c["actual"][:100],
                c["score"],
            ]
            for c in r["error_cases"]
        ]
        _docx_add_table(doc, headers, rows)


def _docx_cost(doc, r):
    items = [
        ["Avg Latency (ms)", r["avg_latency_ms"]],
        ["Min Latency (ms)", r["min_latency_ms"]],
        ["Max Latency (ms)", r["max_latency_ms"]],
        ["First Token (ms)", r["avg_first_token_ms"]],
        [
            "Avg Tokens/Response",
            r["avg_tokens_per_response"],
        ],
        ["Total Tokens", r["total_tokens"]],
        ["Duration (s)", r["duration_seconds"]],
        [
            "Throughput (tok/s)",
            r["throughput_tokens_per_sec"],
        ],
    ]
    _docx_add_table(doc, ["Metric", "Value"], items)


def _docx_value(doc, r):
    doc.add_paragraph(
        f"Score: {r['overall_score']}  |  "
        f"Latency: {r['avg_latency_ms']}ms  |  "
        f"Value Index: {r['value_index']}"
    )
    if r.get("criteria_breakdown"):
        doc.add_heading("Criteria Breakdown", level=2)
        _docx_perf(doc, r)
